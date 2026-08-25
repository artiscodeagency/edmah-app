from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = 'Cahier_des_charges_EDMAH_Django_DRF.docx'
BLUE, DARK, INK = '2E74B5', '1F4D78', '1F2937'
LIGHT_BLUE, LIGHT_GRAY, GOLD, RED = 'E8EEF5', 'F2F4F7', 'FFF4E5', 'FDECEC'

def font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, color):
    pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color); pr.append(shd)

def margins(cell):
    pr = cell._tc.get_or_add_tcPr(); mar = pr.first_child_found_in('w:tcMar')
    if mar is None: mar = OxmlElement('w:tcMar'); pr.append(mar)
    for side, val in [('top', 80), ('start', 120), ('bottom', 80), ('end', 120)]:
        n = mar.find(qn('w:' + side))
        if n is None: n = OxmlElement('w:' + side); mar.append(n)
        n.set(qn('w:w'), str(val)); n.set(qn('w:type'), 'dxa')

def fixed(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    pr = table._tbl.tblPr
    for tag, val in [('tblW', '9360'), ('tblInd', '120')]:
        n = pr.first_child_found_in('w:' + tag)
        if n is None: n = OxmlElement('w:' + tag); pr.append(n)
        n.set(qn('w:w'), val); n.set(qn('w:type'), 'dxa')
    for col, width in zip(table._tbl.tblGrid.gridCol_lst, widths): col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            pr = cell._tc.get_or_add_tcPr(); n = pr.find(qn('w:tcW'))
            if n is None: n = OxmlElement('w:tcW'); pr.append(n)
            n.set(qn('w:w'), str(width)); n.set(qn('w:type'), 'dxa'); margins(cell)

def repeat(row):
    pr = row._tr.get_or_add_trPr(); n = OxmlElement('w:tblHeader'); n.set(qn('w:val'), 'true'); pr.append(n)

def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = 'Table Grid'; fixed(table, widths); repeat(table.rows[0])
    for cell, text in zip(table.rows[0].cells, headers):
        shade(cell, LIGHT_BLUE); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); r = p.add_run(text); font(r, 10, True, DARK)
    for data in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, data):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); r = p.add_run(text); font(r, 9.6)
    return table

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text); font(r, 10.5)

def number(doc, text):
    p = doc.add_paragraph(style='List Number'); p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text); font(r, 10.5)

def para(doc, text, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text); font(r, 10.5); return p

doc = Document(); sec = doc.sections[0]
for a in ['top_margin', 'bottom_margin', 'left_margin', 'right_margin']: setattr(sec, a, Inches(1))
sec.header_distance = Inches(.492); sec.footer_distance = Inches(.492)
normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri'); normal.font.size = Pt(11); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in [('Heading 1',16,BLUE,18,10), ('Heading 2',13,BLUE,14,7), ('Heading 3',12,DARK,10,5)]:
    s = doc.styles[name]; s.font.name = 'Calibri'; s._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); s._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); s.font.size = Pt(size); s.font.bold = True; s.font.color.rgb = RGBColor.from_string(color); s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after); s.paragraph_format.keep_with_next = True

# Cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(92); p.paragraph_format.space_after=Pt(12); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('EDMAH'); font(r,15,True,DARK)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(12); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('CAHIER DES CHARGES'); font(r,30,True,DARK)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(26); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Plateforme web de formation et gestion des inscriptions'); font(r,15,False,'4B5563')
for label,value in [('Projet','EDMAH - École du Mariage Harmonieux'), ('Version','1.0'), ('Date','25 août 2026'), ('Architecture retenue','Django + Django REST Framework')]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(3); r=p.add_run(label+' : '); font(r,10.5,True,DARK); r=p.add_run(value); font(r,10.5)
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(28); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Document de référence pour la conception, le développement, la recette et la mise en production.'); font(r,10,False,'6B7280',True)
doc.add_page_break()

doc.add_paragraph('1. Présentation du projet', style='Heading 1')
para(doc, 'EDMAH souhaite disposer d’une plateforme web complète pour présenter ses activités, gérer les inscriptions aux formations et événements, diffuser des contenus pédagogiques et accompagner les apprenants jusqu’à la certification. Le front-end existant sert de base visuelle et fonctionnelle ; il doit être raccordé à un back-end sécurisé et administrable.')
doc.add_paragraph('1.1 Contexte', style='Heading 2')
para(doc, 'Le site actuel est une maquette front-end composée de pages publiques et d’écrans d’inscription, de cours et de certification. Les données sont statiques et plusieurs actions sont simulées. Le projet vise à rendre ces actions réelles, fiables et administrables.')
doc.add_paragraph('1.2 Objectifs', style='Heading 2')
for x in ['Présenter clairement l’offre EDMAH et faciliter l’acquisition de nouveaux candidats.', 'Centraliser la gestion administrative des formations, événements, contenus et dossiers.', 'Proposer un espace numérique sécurisé pour les apprenants.', 'Assurer la traçabilité des inscriptions, évaluations, progressions et certificats.', 'Permettre à l’équipe EDMAH de gérer le contenu sans modifier le code.']: bullet(doc,x)
doc.add_paragraph('1.3 Périmètre', style='Heading 2')
add_table(doc, ['Inclus', 'Non inclus dans le MVP'], [
    ('Site public, catalogue, blog, galerie, contact, newsletter, inscriptions, administration, espace apprenant et quiz.', 'Application mobile native, visioconférence intégrée, comptabilité complète et marketplace multi-écoles.'),
    ('Gestion des événements, ressources pédagogiques, progression et certification selon règles validées.', 'Forum avancé, paiement automatisé et tableaux de bord avancés : prévus en phase ultérieure si non nécessaires au lancement.'),
], [4680,4680])

doc.add_paragraph('2. Parties prenantes et rôles', style='Heading 1')
add_table(doc, ['Acteur', 'Responsabilités / accès'], [
    ('Direction EDMAH', 'Valide les règles métier, contenus, tarifs, admissions, calendrier, conformité et mise en production.'),
    ('Administrateur', 'Gère les utilisateurs, formations, événements, articles, dossiers, paiements, certificats et paramètres via Django Admin.'),
    ('Formateur', 'Gère ou consulte ses modules, ressources, évaluations et échanges selon les droits accordés.'),
    ('Candidat', 'Dépose un dossier d’inscription, transmet les documents demandés et suit le statut de sa demande.'),
    ('Apprenant', 'Accède aux formations validées, leçons, ressources, quiz, notes, progression et certificat.'),
    ('Visiteur', 'Consulte les pages publiques, recherche des contenus, contacte EDMAH et s’abonne à la newsletter.'),
    ('Équipe technique', 'Conçoit, développe, teste, déploie, sécurise et maintient la plateforme.'),
], [2500,6860])

doc.add_paragraph('3. Exigences fonctionnelles', style='Heading 1')
doc.add_paragraph('3.1 Site public et contenus', style='Heading 2')
add_table(doc, ['Réf.', 'Exigence', 'Critère d’acceptation'], [
    ('F-01', 'Accueil : présentation de l’école, chiffres, formations, témoignages et vidéo.', 'Les contenus sont administrables ; la vidéo et les appels à l’action fonctionnent.'),
    ('F-02', 'Catalogue de formations avec filtres, détail, prix, durée, contenu et modes.', 'Le catalogue est alimenté par l’API et la formation peut être pré-sélectionnée à l’inscription.'),
    ('F-03', 'Événements : publication, filtres, détails, capacité et inscription dédiée.', 'Un administrateur gère les événements ; le nombre de places est contrôlé.'),
    ('F-04', 'Blog : articles, catégories, tags, recherche et pagination.', 'Les articles publiés apparaissent sans modification du code.'),
    ('F-05', 'Galerie : photos et vidéos, filtres et visionneuse.', 'Les médias sont gérés depuis l’administration et la visionneuse fonctionne sur mobile.'),
    ('F-06', 'Contact, FAQ, carte, coordonnées et newsletter.', 'Les demandes et abonnements sont enregistrés et exploitables dans l’administration.'),
], [720,3600,5040])
doc.add_paragraph('3.2 Inscription et dossiers candidats', style='Heading 2')
add_table(doc, ['Réf.', 'Exigence', 'Critère d’acceptation'], [
    ('F-07', 'Formulaire d’inscription en étapes : identité, formation, mode, documents et récapitulatif.', 'Les validations sont faites dans le navigateur et sur le serveur ; le dossier est enregistré.'),
    ('F-08', 'Téléversement de photo, document d’état civil et pièces facultatives.', 'Les types et tailles sont contrôlés ; les fichiers sont privés et accessibles aux personnes autorisées.'),
    ('F-09', 'Statuts de dossier : reçu, en cours, complément demandé, accepté, refusé.', 'Le candidat et l’administration voient un statut clair et l’historique est conservé.'),
    ('F-10', 'Notifications email liées au dossier.', 'Confirmation, demande de complément et décision sont envoyées avec des modèles paramétrables.'),
    ('F-11', 'Inscription aux événements.', 'Le contexte de l’événement est transmis ; les places et inscriptions sont gérées séparément des formations.'),
], [720,3600,5040])
doc.add_paragraph('3.3 Comptes, accès et espace apprenant', style='Heading 2')
add_table(doc, ['Réf.', 'Exigence', 'Critère d’acceptation'], [
    ('F-12', 'Création de compte, connexion, déconnexion, réinitialisation de mot de passe et rôles.', 'Les droits sont contrôlés côté serveur ; un utilisateur ne voit que ses informations autorisées.'),
    ('F-13', 'Tableau de bord apprenant.', 'L’apprenant voit ses formations, sa progression, les actions à effectuer et les dernières informations.'),
    ('F-14', 'Cours organisés en modules et leçons.', 'Les contenus sont affichés dans l’ordre et les règles de déblocage sont appliquées.'),
    ('F-15', 'Ressources pédagogiques protégées.', 'Un apprenant autorisé télécharge une ressource réelle ; un visiteur non autorisé est refusé.'),
    ('F-16', 'Notes personnelles.', 'Les notes sont sauvegardées par utilisateur et peuvent être exportées.'),
], [720,3600,5040])
doc.add_paragraph('3.4 Évaluations, communauté et certificat', style='Heading 2')
add_table(doc, ['Réf.', 'Exigence', 'Critère d’acceptation'], [
    ('F-17', 'Quiz avec questions, réponses, score, seuil de validation et tentatives.', 'Le score est calculé côté serveur et la progression est enregistrée.'),
    ('F-18', 'Forum de discussion modéré (phase ultérieure).', 'Les publications peuvent être créées, signalées et modérées selon les droits.'),
    ('F-19', 'Certification.', 'Le certificat est délivré uniquement lorsque les règles de réussite sont respectées.'),
    ('F-20', 'Vérification publique d’un certificat.', 'Un code unique permet de vérifier l’authenticité et le statut du certificat.'),
], [720,3600,5040])
doc.add_paragraph('3.5 Administration Django', style='Heading 2')
para(doc, 'Django Admin est l’outil d’administration principal. Il doit permettre de gérer au minimum : utilisateurs et rôles, formations, sessions, modules, leçons, ressources, quiz, événements, articles, médias, newsletters, demandes de contact, dossiers d’inscription, documents, certificats et paramètres métier.')
for x in ['Les listes d’administration doivent proposer recherche, filtres, tri et actions groupées lorsque pertinent.', 'Les actions sensibles (acceptation, refus, délivrance de certificat) doivent être réservées aux rôles habilités.', 'Les suppressions de données importantes doivent être limitées et traçables.']: bullet(doc,x)

doc.add_page_break()
doc.add_paragraph('4. Exigences non fonctionnelles', style='Heading 1')
add_table(doc, ['Domaine', 'Exigences'], [
    ('Sécurité', 'HTTPS obligatoire, mots de passe hachés, permissions Django/DRF, protection CSRF lorsque nécessaire, limitation des tentatives, validation serveur, journalisation des opérations sensibles.'),
    ('Données personnelles', 'Accès par rôle, stockage privé des pièces d’identité, durée de conservation définie, politique de confidentialité et possibilité de traitement des demandes de suppression/export.'),
    ('Performance', 'Pages publiques optimisées, pagination des listes, images compressées, chargement différé des médias et API paginées.'),
    ('Compatibilité', 'Utilisation sur mobile, tablette et ordinateur ; compatibilité avec les navigateurs récents.'),
    ('Accessibilité', 'Navigation au clavier, contrastes lisibles, libellés de formulaires, messages d’erreur compréhensibles et alternatives textuelles pour les images importantes.'),
    ('Fiabilité', 'Sauvegardes régulières de la base et des fichiers, gestion des erreurs, surveillance des services et procédure de restauration.'),
    ('Maintenabilité', 'Code organisé par applications Django, documentation technique, conventions de développement, tests automatisés et configuration séparée par environnement.'),
], [2200,7160])

doc.add_paragraph('5. Architecture technique cible', style='Heading 1')
add_table(doc, ['Composant', 'Choix / rôle'], [
    ('Front-end', 'Interface existante adaptée pour consommer les API ; navigation responsive et composants d’état (chargement, succès, erreur).'),
    ('Back-end', 'Django : logique métier, modèles, administration, emails, tâches de gestion et sécurité applicative.'),
    ('API', 'Django REST Framework : sérialiseurs, vues/API, permissions, pagination, filtres et documentation OpenAPI.'),
    ('Authentification', 'Compte Django personnalisé si nécessaire ; authentification sécurisée par session ou jetons selon l’architecture du front-end.'),
    ('Base de données', 'PostgreSQL en production. SQLite est accepté uniquement pour le développement local.'),
    ('Fichiers', 'Stockage objet ou stockage serveur protégé pour documents privés ; médias publics optimisés et servis avec cache.'),
    ('Tâches asynchrones', 'Service de file d’attente si nécessaire pour emails, génération de certificats, exports et traitements de fichiers.'),
    ('Déploiement', 'Environnements développement, recette et production ; serveur applicatif, reverse proxy, HTTPS, sauvegardes et surveillance.'),
], [2200,7160])
doc.add_paragraph('5.1 Applications Django proposées', style='Heading 2')
add_table(doc, ['Application', 'Responsabilité'], [
    ('accounts', 'Utilisateurs, rôles, profils, authentification et droits.'),
    ('catalog', 'Formations, catégories, sessions, tarifs et modalités.'),
    ('admissions', 'Dossiers d’inscription, documents, statuts et décisions.'),
    ('learning', 'Modules, leçons, ressources, progression, notes et quiz.'),
    ('events', 'Événements, inscriptions, capacité et informations pratiques.'),
    ('content', 'Blog, pages institutionnelles, galerie, FAQ et médias.'),
    ('communications', 'Contact, newsletter, modèles email et notifications.'),
    ('certificates', 'Règles de délivrance, documents de certificat et vérification.'),
], [2200,7160])

doc.add_paragraph('6. Données et API', style='Heading 1')
doc.add_paragraph('6.1 Données principales', style='Heading 2')
para(doc, 'Les principales entités sont : utilisateur, rôle, profil, formation, session, module, leçon, ressource, inscription formation, document, événement, inscription événement, article, catégorie, tag, média galerie, message contact, abonnement newsletter, quiz, question, tentative, progression, note, publication forum, certificat et paiement.')
doc.add_paragraph('6.2 Principes API', style='Heading 2')
for x in ['API versionnées et documentées ; format JSON pour les données métier et multipart/form-data pour les téléversements.', 'Validation des entrées par sérialiseurs DRF et règles métier dans les services/modèles Django.', 'Pagination, filtres et recherche pour les listes de formations, événements, articles, inscriptions et utilisateurs.', 'Réponses d’erreur cohérentes, sans exposition de données sensibles.', 'Permissions explicites pour chaque endpoint ; aucune protection ne doit dépendre uniquement du front-end.']: bullet(doc,x)
doc.add_paragraph('6.3 Exemples d’API à fournir', style='Heading 2')
add_table(doc, ['Domaine', 'Exemples d’actions API'], [
    ('Public', 'Lister et consulter formations, événements, articles, catégories, galerie ; envoyer contact ; souscrire newsletter.'),
    ('Admission', 'Créer un dossier, téléverser documents, consulter son dossier, traiter un dossier côté administration.'),
    ('Apprentissage', 'Lister mes formations, consulter leçons, enregistrer progression, soumettre un quiz, gérer notes et ressources.'),
    ('Administration', 'Créer et modifier contenus via Django Admin ; endpoints sécurisés complémentaires si le front-end admin le requiert.'),
    ('Certification', 'Consulter mon certificat et vérifier publiquement un code de certificat.'),
], [2200,7160])

doc.add_page_break()
doc.add_paragraph('7. Sécurité, conformité et sauvegarde', style='Heading 1')
doc.add_paragraph('7.1 Mesures obligatoires', style='Heading 2')
for x in ['Activer HTTPS en production et ne jamais exposer les clés secrètes dans le code ou le dépôt.', 'Contrôler strictement les types, tailles et contenus des fichiers déposés ; prévoir une analyse antivirus si le contexte l’exige.', 'Conserver les pièces d’identité dans un emplacement privé, avec URL temporaire ou contrôle d’accès serveur.', 'Appliquer le principe du moindre privilège aux rôles administratifs et conserver les traces des décisions de dossier.', 'Mettre en place une politique de sauvegarde testée pour la base PostgreSQL et les fichiers.', 'Préparer les mentions légales, conditions d’utilisation, politique de confidentialité et consentement newsletter avant lancement.']: bullet(doc,x)
doc.add_paragraph('7.2 Risques principaux', style='Heading 2')
add_table(doc, ['Risque', 'Mesure de réduction'], [
    ('Exposition de documents personnels', 'Stockage privé, permissions, liens temporaires, audit des accès et durée de conservation encadrée.'),
    ('Règles pédagogiques non définies', 'Valider avant développement les conditions de déblocage, scores, tentatives et certification.'),
    ('Contenus indisponibles ou non validés', 'Désigner un responsable de contenu et constituer un lot initial avant recette.'),
    ('Paiement non choisi', 'Prévoir une phase de décision ; démarrer au besoin par un suivi manuel des paiements.'),
    ('Dépendance aux services externes', 'Documenter les comptes, clés, sauvegardes et procédures de remplacement.'),
], [2600,6760])

doc.add_paragraph('8. Méthode projet, planning et validation', style='Heading 1')
doc.add_paragraph('8.1 Méthode de travail', style='Heading 2')
para(doc, 'Le projet est mené par itérations courtes. Chaque phase se termine par une démonstration, une validation métier et la correction des anomalies bloquantes. Les choix fonctionnels non définis font l’objet d’une décision formalisée avant le développement de la phase concernée.')
doc.add_paragraph('8.2 Planning indicatif', style='Heading 2')
add_table(doc, ['Phase', 'Durée', 'Livrable de fin de phase'], [
    ('Cadrage', 'Semaine 1', 'Règles métier, contenus initiaux, critères de réussite et choix des services validés.'),
    ('Fondations', 'Semaines 2 à 3', 'Projet Django/DRF, base, rôles, sécurité initiale et administration prête.'),
    ('Site public', 'Semaines 4 à 5', 'Catalogue, événements, blog, galerie, contact et newsletter alimentés par API.'),
    ('Admissions', 'Semaines 6 à 7', 'Dossiers, documents, statuts, traitement administratif et emails opérationnels.'),
    ('Espace apprenant', 'Semaines 8 à 10', 'Cours, progression, ressources, notes et quiz sur une première formation complète.'),
    ('Fonctions avancées', 'Semaines 11 à 12', 'Paiement, forum, certificats et tableaux de bord selon priorités confirmées.'),
    ('Recette / production', 'Semaines 13 à 14', 'Tests, validation métier, formation utilisateurs et mise en ligne.'),
], [2400,1900,5060])
doc.add_paragraph('8.3 Critères de recette', style='Heading 2')
for x in ['Chaque exigence fonctionnelle prioritaire possède un scénario de test validé.', 'Les liens et appels à l’action mènent vers un parcours réel, sans simulation.', 'Les utilisateurs ne peuvent accéder qu’aux informations prévues pour leur rôle.', 'Les formulaires gèrent clairement les succès et erreurs ; les données sont persistées.', 'Les parcours sont testés sur mobile et ordinateur.', 'Les sauvegardes et la restauration sont vérifiées avant la mise en production.']: bullet(doc,x)

doc.add_paragraph('9. Livrables attendus', style='Heading 1')
for x in ['Code source du back-end Django et des adaptations front-end, organisé et documenté.', 'Schéma de base de données et migrations Django.', 'API REST documentée (OpenAPI/Swagger ou équivalent).', 'Django Admin configuré avec les modèles, filtres, recherches et droits nécessaires.', 'Jeux de données de démonstration ou procédure de saisie des contenus initiaux.', 'Tests automatisés essentiels et compte-rendu de recette.', 'Guide d’installation/déploiement, sauvegarde et exploitation.', 'Guide utilisateur administrateur et guide de gestion des dossiers.', 'Accès aux environnements, variables configurées de manière sécurisée et procédure de mise en production.']: bullet(doc,x)

doc.add_paragraph('10. Hypothèses et décisions à valider', style='Heading 1')
for x in ['Le contenu des formations, sessions, documents et règles pédagogiques sera fourni et validé par EDMAH.', 'Les responsables des admissions, contenus, forum et certificats seront désignés.', 'Le moyen de paiement et les obligations associées seront choisis avant son intégration.', 'Les mentions légales et règles de conservation des dossiers seront validées avant ouverture publique.', 'Le nom de domaine, l’hébergement et les comptes de services externes seront disponibles à temps.']: number(doc,x)

doc.add_paragraph('11. Approbation', style='Heading 1')
para(doc, 'Ce cahier des charges sert de référence de réalisation. Toute évolution qui modifie le périmètre, les délais, les règles métier ou les services externes devra être formalisée et validée avant développement.')
table=add_table(doc,['Partie','Nom / fonction','Date et signature'], [('EDMAH','', ''), ('Prestataire technique','', '')], [2500,3500,3360])
for row in table.rows[1:]:
    for cell in row.cells: cell.paragraphs[0].paragraph_format.space_after=Pt(26)

footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=footer.add_run('EDMAH - Cahier des charges'); font(r,8.5,False,'6B7280')
doc.core_properties.title='Cahier des charges EDMAH - Django et Django REST Framework'; doc.core_properties.subject='Cahier des charges fonctionnel et technique'; doc.core_properties.author='EDMAH'
doc.save(OUT); print(OUT)
