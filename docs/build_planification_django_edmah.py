from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = 'Planification_developpement_EDMAH_Django.docx'
BLUE, DARK, INK = '2E74B5', '1F4D78', '1F2937'
LIGHT_BLUE, LIGHT_GRAY, GOLD = 'E8EEF5', 'F2F4F7', 'FFF4E5'

phases = [
    ('Phase 0', 'Semaine 1', 'Cadrage et préparation', 'P0', [
        ('Valider les règles métier', 'Confirmer les formations, prix, modes, pièces demandées, règles de validation et critères du certificat.'),
        ('Préparer le projet Django', 'Créer le projet, les applications Django, les environnements de configuration et le dépôt de code.'),
        ('Choisir les services externes', 'Valider la base PostgreSQL, le stockage sécurisé des fichiers, le service email et le paiement à intégrer plus tard.'),
    ]),
    ('Phase 1', 'Semaines 2 à 3', 'Fondations techniques et administration', 'P0', [
        ('Configurer la base de données', 'Mettre en place PostgreSQL et les modèles de base : utilisateur, rôle, formation, session et inscription.'),
        ('Créer les comptes et rôles', 'Créer les rôles visiteur, candidat, apprenant, formateur et administrateur avec leurs droits.'),
        ('Configurer Django Admin', 'Rendre les objets importants gérables dans l’administration : formations, événements, articles, utilisateurs et inscriptions.'),
        ('Créer les API REST de base', 'Mettre en place Django REST Framework, les sérialiseurs, permissions, pagination et documentation des API.'),
        ('Sécuriser le projet', 'Ajouter authentification, validation serveur, variables secrètes, protection des fichiers et journalisation minimale.'),
    ]),
    ('Phase 2', 'Semaines 4 à 5', 'Site public et gestion des contenus', 'P0', [
        ('API des formations', 'Fournir le catalogue, les catégories, les détails, les sessions, les prix et les filtres.'),
        ('API des événements', 'Publier les événements, leurs dates, places, types et informations pratiques.'),
        ('API blog et galerie', 'Permettre de publier articles, catégories, tags, images et vidéos depuis l’administration.'),
        ('Raccorder le front-end', 'Remplacer les données écrites dans les pages par les API Django et corriger les liens d’inscription existants.'),
        ('Contact et newsletter', 'Enregistrer les messages et abonnements, vérifier les emails et permettre une gestion dans Django Admin.'),
    ]),
    ('Phase 3', 'Semaines 6 à 7', 'Inscription aux formations', 'P0', [
        ('API du formulaire multi-étapes', 'Enregistrer identité, formation choisie, mode de formation, motivations et acceptation des conditions.'),
        ('Gestion des documents', 'Téléverser les pièces justificatives avec contrôle de format, taille, accès privé et traçabilité.'),
        ('Traitement des dossiers', 'Donner à l’administration des statuts clairs : reçu, en cours, complément demandé, accepté ou refusé.'),
        ('Emails automatiques', 'Envoyer confirmation de dépôt, demande de complément et décision au candidat.'),
        ('Inscription événement', 'Créer un parcours distinct, lié à un événement, avec contrôle du nombre de places.'),
    ]),
    ('Phase 4', 'Semaines 8 à 10', 'Espace apprenant et cours', 'P1', [
        ('Tableau de bord apprenant', 'Afficher les formations suivies, les modules disponibles, la progression et les actions à faire.'),
        ('Modules et leçons', 'Créer les contenus de cours, organiser leur ordre et contrôler le déblocage des modules.'),
        ('Ressources protégées', 'Donner accès aux documents et vidéos uniquement aux apprenants autorisés.'),
        ('Quiz et progression', 'Enregistrer réponses, scores, tentatives et validation des modules selon le seuil défini.'),
        ('Notes personnelles', 'Sauvegarder les notes privées de l’apprenant et permettre leur export.'),
    ]),
    ('Phase 5', 'Semaines 11 à 12', 'Fonctions avancées', 'P2', [
        ('Paiement', 'Intégrer le moyen retenu : paiement manuel, Mobile Money, carte ou virement, puis suivre son état.'),
        ('Certificats', 'Produire un certificat après contrôle des prérequis et proposer une page de vérification.'),
        ('Forum et modération', 'Permettre les échanges, la signalisation et la modération par les administrateurs ou formateurs.'),
        ('Statistiques', 'Afficher les indicateurs utiles : inscriptions, dossiers, progression, résultats et événements.'),
    ]),
    ('Phase 6', 'Semaines 13 à 14', 'Recette et mise en ligne', 'P0', [
        ('Tests complets', 'Tester les parcours visiteur, candidat, apprenant, formateur et administrateur sur ordinateur et mobile.'),
        ('Recette métier', 'Faire valider les écrans et règles par l’équipe EDMAH, puis corriger les anomalies prioritaires.'),
        ('Préparer la production', 'Configurer le domaine, HTTPS, sauvegardes, surveillance, emails de production et accès administrateurs.'),
        ('Former l’équipe', 'Montrer comment gérer les contenus, dossiers, événements, utilisateurs et demandes reçues.'),
    ]),
]

def set_font(run, size=11, bold=False, color=INK):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color); tc_pr.append(shd)

def cell_margin(cell):
    tc_pr = cell._tc.get_or_add_tcPr(); tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None: tc_mar = OxmlElement('w:tcMar'); tc_pr.append(tc_mar)
    for side, value in [('top',80), ('start',120), ('bottom',80), ('end',120)]:
        n = tc_mar.find(qn('w:'+side))
        if n is None: n = OxmlElement('w:'+side); tc_mar.append(n)
        n.set(qn('w:w'), str(value)); n.set(qn('w:type'),'dxa')

def fixed(table, widths):
    table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.LEFT
    pr=table._tbl.tblPr
    for tag,val in [('tblW','9360'),('tblInd','120')]:
        n=pr.first_child_found_in('w:'+tag)
        if n is None: n=OxmlElement('w:'+tag); pr.append(n)
        n.set(qn('w:w'),val); n.set(qn('w:type'),'dxa')
    for col,width in zip(table._tbl.tblGrid.gridCol_lst,widths): col.set(qn('w:w'),str(width))
    for row in table.rows:
        for cell,width in zip(row.cells,widths):
            tcpr=cell._tc.get_or_add_tcPr(); n=tcpr.find(qn('w:tcW'))
            if n is None: n=OxmlElement('w:tcW'); tcpr.append(n)
            n.set(qn('w:w'),str(width)); n.set(qn('w:type'),'dxa'); cell_margin(cell)

def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr(); n = OxmlElement('w:tblHeader'); n.set(qn('w:val'),'true'); tr_pr.append(n)

doc=Document(); sec=doc.sections[0]
for attr in ['top_margin','bottom_margin','left_margin','right_margin']: setattr(sec,attr,Inches(1))
sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)
normal=doc.styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for style_name,size,color,before,after in [('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,14,7),('Heading 3',12,DARK,10,5)]:
    s=doc.styles[style_name]; s.font.name='Calibri'; s._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); s._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

# Editorial first page header
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(28); p.paragraph_format.space_after=Pt(8); r=p.add_run('EDMAH'); set_font(r,12,True,DARK)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(8); r=p.add_run('Planification du développement'); set_font(r,28,True,DARK)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(20); r=p.add_run('Architecture Django + Django REST Framework'); set_font(r,13,False,'4B5563')
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(12); r=p.add_run('Objectif : '); set_font(r,11,True,DARK); r=p.add_run('livrer progressivement un site public, une gestion administrative et un espace apprenant sécurisé.'); set_font(r)

doc.add_paragraph('Choix technique retenu', style='Heading 1')
for label,text in [
    ('Back-end','Django pour la logique métier, la gestion des utilisateurs et l’interface d’administration.'),
    ('API','Django REST Framework pour fournir les API utilisées par le front-end.'),
    ('Base de données','PostgreSQL pour les données de production.'),
    ('Administration','Django Admin pour gérer rapidement les formations, contenus, inscriptions et utilisateurs.'),
    ('Fichiers','Stockage privé pour les documents d’inscription et stockage adapté pour les médias publics.'),
]:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); r=p.add_run(label+' : '); set_font(r,10.5,True,DARK); r=p.add_run(text); set_font(r,10.5)

doc.add_paragraph('Plan par phases', style='Heading 1')
for phase, period, title, priority, tasks in phases:
    doc.add_paragraph(f'{phase} - {title}', style='Heading 2')
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); r=p.add_run(f'Période : {period} | Priorité : {priority}'); set_font(r,10,True, DARK)
    table=doc.add_table(rows=1, cols=2); table.style='Table Grid'; fixed(table,[2850,6510]); repeat_header(table.rows[0])
    for cell,text in zip(table.rows[0].cells,['Tâche','Résultat attendu']):
        shade(cell,LIGHT_BLUE); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(text); set_font(r,10.5,True,DARK)
    for task,result in tasks:
        cells=table.add_row().cells
        for c in cells: c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
        p=cells[0].paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(task); set_font(r,10.2,True,DARK)
        p=cells[1].paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(result); set_font(r,10.2)

doc.add_paragraph('Règles de réalisation', style='Heading 1')
rules=[
    'Commencer par les tâches P0 : elles sont nécessaires avant toute ouverture au public.',
    'Chaque API doit être testée, documentée et protégée par les permissions Django appropriées.',
    'Les documents personnels déposés par les candidats ne doivent jamais être accessibles publiquement.',
    'Le front-end est raccordé phase par phase : aucune simulation ne doit rester sur une fonction livrée.',
    'À la fin de chaque phase, une démonstration et une validation métier sont prévues avant de poursuivre.',
]
for text in rules:
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(4); r=p.add_run(text); set_font(r,10.5)

doc.add_paragraph('Livrables attendus à la fin du projet', style='Heading 1')
for text in [
    'Une administration Django utilisable par l’équipe EDMAH.',
    'Des API Django REST Framework documentées et utilisées par le site.',
    'Un parcours d’inscription sécurisé et traçable.',
    'Un espace apprenant fonctionnel pour au moins une formation complète.',
    'Une version prête à être mise en ligne, avec sauvegarde, HTTPS et accès administrateurs configurés.',
]:
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(4); r=p.add_run(text); set_font(r,10.5)

footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=footer.add_run('EDMAH - Planification Django + DRF'); set_font(r,8.5,False,'6B7280')
doc.core_properties.title='Planification du développement EDMAH - Django et DRF'; doc.core_properties.subject='Plan de développement'; doc.core_properties.author='EDMAH'
doc.save(OUT); print(OUT)
