from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = 'Fonctionnalites_EDMAH_priorisees.docx'
BLUE = '2E74B5'
DARK = '1F4D78'
LIGHT = 'E8EEF5'
GRAY = 'F2F4F7'
INK = '1F2937'

items = [
    ('1', 'P0 - Indispensable', 'Inscription aux formations', 'Permettre à un candidat de remplir son dossier, choisir une formation et envoyer sa demande.'),
    ('2', 'P0 - Indispensable', 'Gestion des dossiers d’inscription', 'Permettre à l’administration de consulter les dossiers, demander un complément, accepter ou refuser une demande.'),
    ('3', 'P0 - Indispensable', 'Comptes utilisateurs et connexion', 'Donner à chaque candidat, apprenant, formateur et administrateur un accès sécurisé adapté à son rôle.'),
    ('4', 'P0 - Indispensable', 'Gestion des formations', 'Créer, modifier, publier ou retirer les formations, leurs prix, durées, sessions et modalités.'),
    ('5', 'P0 - Indispensable', 'Dépôt sécurisé des documents', 'Recevoir les pièces demandées (photo, état civil, PDF) avec contrôle de format et de taille.'),
    ('6', 'P0 - Indispensable', 'Notifications par email', 'Confirmer une inscription, informer sur l’état du dossier et envoyer les consignes utiles.'),
    ('7', 'P0 - Indispensable', 'Formulaire de contact', 'Enregistrer les messages reçus et permettre à l’équipe EDMAH de les traiter.'),
    ('8', 'P0 - Indispensable', 'Newsletter', 'Collecter les abonnements avec consentement et permettre l’envoi d’informations aux abonnés.'),
    ('9', 'P1 - Important', 'Espace apprenant', 'Donner aux apprenants inscrits un tableau de bord avec leurs formations et leur progression.'),
    ('10', 'P1 - Important', 'Cours organisés par modules', 'Afficher les leçons dans le bon ordre et débloquer les modules selon les règles de formation.'),
    ('11', 'P1 - Important', 'Quiz et évaluations', 'Enregistrer les réponses, calculer le score et valider un module lorsque le seuil demandé est atteint.'),
    ('12', 'P1 - Important', 'Ressources téléchargeables', 'Mettre à disposition les supports de cours uniquement aux apprenants autorisés.'),
    ('13', 'P1 - Important', 'Notes personnelles', 'Permettre à chaque apprenant d’écrire, sauvegarder et exporter ses notes de cours.'),
    ('14', 'P1 - Important', 'Catalogue et filtres de formations', 'Afficher les formations à partir de vraies données, avec recherche et filtres par catégorie.'),
    ('15', 'P1 - Important', 'Gestion des événements', 'Publier les événements, gérer les places et permettre une inscription spécifique à chaque événement.'),
    ('16', 'P1 - Important', 'Blog, galerie et contenus publics', 'Permettre à l’équipe de publier des articles, photos et vidéos sans modifier le code du site.'),
    ('17', 'P2 - À planifier ensuite', 'Paiement des formations et événements', 'Accepter et suivre les paiements selon la solution choisie : Mobile Money, carte, virement ou paiement manuel.'),
    ('18', 'P2 - À planifier ensuite', 'Forum de discussion', 'Permettre les échanges entre apprenants, avec modération par l’équipe EDMAH.'),
    ('19', 'P2 - À planifier ensuite', 'Certificats', 'Générer un certificat après validation des conditions de formation et proposer une vérification du certificat.'),
    ('20', 'P2 - À planifier ensuite', 'Statistiques et tableaux de bord', 'Suivre les inscriptions, la progression, les événements, les résultats et les abonnements.'),
]

def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')

def set_fixed_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), '9360')
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.first_child_found_in('w:tblInd')
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margin(cell)

def set_run_font(run, name='Calibri', size=11, bold=False, color=INK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ('Heading 1', 16, BLUE, 18, 10),
    ('Heading 2', 13, BLUE, 14, 7),
    ('Heading 3', 12, DARK, 10, 5),
]:
    style = styles[name]
    style.font.name = 'Calibri'
    style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

# Simple editorial cover/title block
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(32)
p.paragraph_format.space_after = Pt(8)
r = p.add_run('EDMAH')
set_run_font(r, size=12, bold=True, color=DARK)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run('Fonctionnalités à réaliser')
set_run_font(r, size=28, bold=True, color=DARK)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(20)
r = p.add_run('Liste simple et priorisée, du plus important au moins important')
set_run_font(r, size=13, color='4B5563')

intro = doc.add_paragraph()
intro.paragraph_format.space_after = Pt(14)
r = intro.add_run('Lecture des priorités : ')
set_run_font(r, bold=True, color=DARK)
r = intro.add_run('P0 = indispensable pour démarrer, P1 = important pour une bonne première version, P2 = à planifier après le lancement.')
set_run_font(r, color=INK)

doc.add_paragraph('Fonctionnalités classées par priorité', style='Heading 1')

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.style = 'Table Grid'
set_fixed_table_geometry(table, [540, 2700, 6120])
headers = ['N°', 'Fonctionnalité', 'Commentaire simple']
for cell, text in zip(table.rows[0].cells, headers):
    set_cell_shading(cell, LIGHT)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=True, color=DARK)

for number, priority, feature, comment in items:
    row = table.add_row()
    row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    row.cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if priority.startswith('P0'):
        fill = 'FFF4E5'
    elif priority.startswith('P1'):
        fill = 'F2F7FC'
    else:
        fill = GRAY
    set_cell_shading(row.cells[0], fill)
    p = row.cells[0].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(number)
    set_run_font(r, size=10.5, bold=True, color=DARK)
    p = row.cells[1].paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(feature)
    set_run_font(r, size=10.5, bold=True, color=DARK)
    p = row.cells[1].add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(priority)
    set_run_font(r, size=8.5, color='6B7280')
    p = row.cells[2].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(comment)
    set_run_font(r, size=10.5, color=INK)

doc.add_paragraph('Ordre recommandé de démarrage', style='Heading 1')
for text in [
    'Mettre en place les comptes, la gestion des formations et les dossiers d’inscription.',
    'Brancher les emails, le contact et la newsletter.',
    'Ouvrir ensuite l’espace apprenant avec une première formation complète.',
    'Ajouter enfin les paiements, le forum, les certificats et les statistiques.',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=10.5)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.paragraph_format.space_before = Pt(4)
r = footer.add_run('EDMAH - Priorités fonctionnelles')
set_run_font(r, size=8.5, color='6B7280')

doc.core_properties.title = 'Fonctionnalités EDMAH priorisées'
doc.core_properties.subject = 'Fonctionnalités à réaliser pour EDMAH'
doc.core_properties.author = 'EDMAH'
doc.save(OUT)
print(OUT)
